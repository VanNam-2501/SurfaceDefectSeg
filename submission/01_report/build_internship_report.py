from __future__ import annotations

import csv
import json
import math
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_COLOR_INDEX, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
REPORT_DIR = ROOT / "submission" / "01_report"
WORK_DIR = REPORT_DIR / "_report_work"
OUTPUT = REPORT_DIR / "Bao_cao_Thuc_tap_tot_nghiep_Phan_doan_khuyet_tat_be_mat_nhom.docx"
FIG_DIR = ROOT / "artifacts" / "reports" / "final" / "visualizations" / "figures" / "png"
THESIS_DIR = ROOT / "artifacts" / "reports" / "final" / "thesis_evaluation_report"
DECISION_DIR = ROOT / "artifacts" / "reports" / "final" / "decision_and_test_audit"
TRAINING_ROOT = ROOT / "artifacts" / "training_download"

NAVY = "17365D"
BLUE = "2E5F8A"
LIGHT_BLUE = "EAF1F7"
PALE_BLUE = "F4F7FA"
GOLD = "B8862B"
LIGHT_GOLD = "FFF4D6"
INK = "1F2933"
MUTED = "5B6770"
LIGHT_GRAY = "F2F4F7"
BORDER = "C9D3DD"
WHITE = "FFFFFF"
RED = "9B1C1C"
GREEN = "216E5B"


def load_csv(path: Path) -> list[dict[str, str]]:
    with path.open("r", encoding="utf-8-sig", newline="") as handle:
        return list(csv.DictReader(handle))


def load_json(path: Path):
    with path.open("r", encoding="utf-8") as handle:
        return json.load(handle)


def training_run(model: str, run_dir: Path, label: str) -> dict:
    """Load reproducibility files exported by a Kaggle training run."""
    root = TRAINING_ROOT / run_dir
    summary_path = root / "training_summary.json"
    history_path = root / "training_history.csv"
    curve_path = root / "curves" / "learning_curve.png"
    if not summary_path.exists() or not history_path.exists():
        return {"model": model, "label": label, "summary": {}, "history_rows": 0, "curve": curve_path, "path": str(root.relative_to(ROOT)).replace("\\", "/")}
    return {
        "model": model,
        "label": label,
        "summary": load_json(summary_path),
        "history_rows": len(load_csv(history_path)),
        "curve": curve_path,
        "path": str(root.relative_to(ROOT)).replace("\\", "/"),
    }


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=110, bottom=90, end=110) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_borders(table, color=BORDER, size=5) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), str(size))
        tag.set(qn("w:color"), color)


def set_table_geometry(table, widths_cm: list[float], indent_dxa=120) -> None:
    total_dxa = sum(int(round(cm / 2.54 * 1440)) for cm in widths_cm)
    grid_dxa = [int(round(cm / 2.54 * 1440)) for cm in widths_cm]
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total_dxa))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in grid_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Cm(widths_cm[idx])
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(grid_dxa[idx]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def set_font(run, size=None, bold=None, italic=None, color=INK, name="Times New Roman") -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_paragraph_shading(paragraph, fill: str, border: str | None = None) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)
    if border:
        p_bdr = OxmlElement("w:pBdr")
        left = OxmlElement("w:left")
        left.set(qn("w:val"), "single")
        left.set(qn("w:sz"), "18")
        left.set(qn("w:space"), "6")
        left.set(qn("w:color"), border)
        p_bdr.append(left)
        p_pr.append(p_bdr)


def add_field(paragraph, instruction: str, display: str = "") -> None:
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    begin_run = OxmlElement("w:r")
    begin_run.append(begin)
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    instr_run = OxmlElement("w:r")
    instr_run.append(instr)
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    separate_run = OxmlElement("w:r")
    separate_run.append(separate)
    text_run = OxmlElement("w:r")
    text = OxmlElement("w:t")
    text.text = display
    text_run.append(text)
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    end_run = OxmlElement("w:r")
    end_run.append(end)
    paragraph._p.extend([begin_run, instr_run, separate_run, text_run, end_run])


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.0)
    section.header_distance = Cm(1.25)
    section.footer_distance = Cm(1.25)
    section.different_first_page_header_footer = True

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(13)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.first_line_indent = Cm(1.0)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.3

    heading_tokens = {
        "Heading 1": (16, NAVY, 18, 10),
        "Heading 2": (14, BLUE, 14, 7),
        "Heading 3": (13, NAVY, 10, 5),
    }
    for style_name, (size, color, before, after) in heading_tokens.items():
        style = styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.first_line_indent = Cm(0)
        style.paragraph_format.line_spacing = 1.1

    for style_name, left, hanging in (("List Bullet", 1.0, 0.5), ("List Number", 1.0, 0.5)):
        style = styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
        style.font.size = Pt(13)
        style.paragraph_format.left_indent = Cm(left)
        style.paragraph_format.first_line_indent = Cm(-hanging)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    for style_name, size, left in (("TOC 1", 11.0, 0.0), ("TOC 2", 10.5, 0.5), ("TOC 3", 10.0, 1.0)):
        if style_name not in styles:
            style = styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
        else:
            style = styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(INK)
        style.paragraph_format.left_indent = Cm(left)
        style.paragraph_format.first_line_indent = Cm(0)
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(1)
        style.paragraph_format.line_spacing = 1.0

    if "Figure Caption" not in styles:
        cap = styles.add_style("Figure Caption", WD_STYLE_TYPE.PARAGRAPH)
    else:
        cap = styles["Figure Caption"]
    cap.font.name = "Times New Roman"
    cap._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    cap._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    cap.font.size = Pt(11)
    cap.font.italic = True
    cap.font.color.rgb = RGBColor.from_string(MUTED)
    cap.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_before = Pt(4)
    cap.paragraph_format.space_after = Pt(10)
    cap.paragraph_format.keep_with_next = False

    if "Table Caption" not in styles:
        tcap = styles.add_style("Table Caption", WD_STYLE_TYPE.PARAGRAPH)
    else:
        tcap = styles["Table Caption"]
    tcap.font.name = "Times New Roman"
    tcap._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    tcap._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    tcap.font.size = Pt(11)
    tcap.font.italic = True
    tcap.font.color.rgb = RGBColor.from_string(MUTED)
    tcap.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tcap.paragraph_format.space_before = Pt(8)
    tcap.paragraph_format.space_after = Pt(5)
    tcap.paragraph_format.keep_with_next = True

    settings = doc.settings._element
    update_fields = OxmlElement("w:updateFields")
    update_fields.set(qn("w:val"), "true")
    settings.append(update_fields)


def set_running_furniture(section) -> None:
    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("BÁO CÁO THỰC TẬP TỐT NGHIỆP  |  PHÂN ĐOẠN KHUYẾT TẬT BỀ MẶT NHÔM")
    set_font(r, size=9, bold=True, color=MUTED)
    p_pr = p._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "5")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), BORDER)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)

    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    r = p.add_run("Trang ")
    set_font(r, size=10, color=MUTED)
    add_field(p, "PAGE", "1")


def add_cover(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("[ĐIỀN TÊN TRƯỜNG]")
    set_font(r, size=14, bold=True, color=NAVY)
    r.font.highlight_color = WD_COLOR_INDEX.YELLOW
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(42)
    r = p.add_run("[ĐIỀN TÊN KHOA/BỘ MÔN]")
    set_font(r, size=13, bold=True, color=NAVY)
    r.font.highlight_color = WD_COLOR_INDEX.YELLOW

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run("BÁO CÁO THỰC TẬP TỐT NGHIỆP")
    set_font(r, size=18, bold=True, color=GOLD)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)
    p.paragraph_format.line_spacing = 1.15
    r = p.add_run(
        "XÂY DỰNG HỆ THỐNG PHÂN ĐOẠN VÀ HỖ TRỢ QUYẾT ĐỊNH\n"
        "PHÁT HIỆN KHUYẾT TẬT BỀ MẶT NHÔM DỰA TRÊN HỌC SÂU"
    )
    set_font(r, size=22, bold=True, color=NAVY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(46)
    r = p.add_run("Development of a Deep Learning-Based Aluminum Surface Defect Segmentation and Decision Support System")
    set_font(r, size=12, italic=True, color=MUTED)

    items = [
        ("Sinh viên thực hiện", "[ĐIỀN HỌ VÀ TÊN]"),
        ("Mã số sinh viên", "[ĐIỀN MSSV]"),
        ("Lớp / Khóa", "[ĐIỀN LỚP / KHÓA]"),
        ("Đơn vị thực tập", "[ĐIỀN TÊN ĐƠN VỊ]"),
        ("Người hướng dẫn tại đơn vị", "[ĐIỀN HỌ TÊN]"),
        ("Giảng viên hướng dẫn", "[ĐIỀN HỌ TÊN]"),
    ]
    for label, value in items:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(3.0)
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(f"{label}: ")
        set_font(r, size=13, bold=True, color=INK)
        r = p.add_run(value)
        set_font(r, size=13, color=INK)
        r.font.highlight_color = WD_COLOR_INDEX.YELLOW

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(56)
    r = p.add_run("TP. HỒ CHÍ MINH, 2026")
    set_font(r, size=13, bold=True, color=NAVY)


def add_page_title(doc: Document, title: str) -> None:
    p = doc.add_paragraph()
    p.style = doc.styles["Heading 1"]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(14)
    r = p.add_run(title)
    set_font(r, size=16, bold=True, color=NAVY)


def add_para(doc: Document, text: str, bold_lead: str | None = None, italic=False, align=None, indent=True):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    if not indent:
        p.paragraph_format.first_line_indent = Cm(0)
    if bold_lead and text.startswith(bold_lead):
        r = p.add_run(bold_lead)
        set_font(r, size=13, bold=True, color=INK)
        r = p.add_run(text[len(bold_lead):])
        set_font(r, size=13, italic=italic, color=INK)
    else:
        r = p.add_run(text)
        set_font(r, size=13, italic=italic, color=INK)
    return p


def new_numbering_instance(doc: Document) -> int:
    style_num_pr = doc.styles["List Number"]._element.pPr.numPr
    base_num_id = int(style_num_pr.numId.val)
    numbering = doc.part.numbering_part.element
    base_num = next(node for node in numbering.findall(qn("w:num")) if int(node.get(qn("w:numId"))) == base_num_id)
    abstract_id = base_num.find(qn("w:abstractNumId")).get(qn("w:val"))
    new_num_id = max(int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))) + 1
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(new_num_id))
    abstract = OxmlElement("w:abstractNumId")
    abstract.set(qn("w:val"), abstract_id)
    num.append(abstract)
    override = OxmlElement("w:lvlOverride")
    override.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:startOverride")
    start.set(qn("w:val"), "1")
    override.append(start)
    num.append(override)
    numbering.append(num)
    return new_num_id


def add_bullets(doc: Document, items: list[str], numbered=False) -> None:
    style = "List Number" if numbered else "List Bullet"
    num_id = new_numbering_instance(doc) if numbered else None
    for item in items:
        p = doc.add_paragraph(style=style)
        p.paragraph_format.first_line_indent = Cm(-0.5)
        if num_id is not None:
            p_pr = p._p.get_or_add_pPr()
            num_pr = p_pr.find(qn("w:numPr"))
            if num_pr is None:
                num_pr = OxmlElement("w:numPr")
                p_pr.append(num_pr)
            ilvl = num_pr.find(qn("w:ilvl"))
            if ilvl is None:
                ilvl = OxmlElement("w:ilvl")
                num_pr.append(ilvl)
            ilvl.set(qn("w:val"), "0")
            num_id_node = num_pr.find(qn("w:numId"))
            if num_id_node is None:
                num_id_node = OxmlElement("w:numId")
                num_pr.append(num_id_node)
            num_id_node.set(qn("w:val"), str(num_id))
        r = p.add_run(item)
        set_font(r, size=13, color=INK)


def add_callout(doc: Document, label: str, text: str, kind="info") -> None:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.left_indent = Cm(0.3)
    p.paragraph_format.right_indent = Cm(0.25)
    p.paragraph_format.space_before = Pt(7)
    p.paragraph_format.space_after = Pt(9)
    p.paragraph_format.line_spacing = 1.2
    fill = LIGHT_GOLD if kind == "caution" else LIGHT_BLUE
    border = GOLD if kind == "caution" else BLUE
    set_paragraph_shading(p, fill, border)
    r = p.add_run(f"{label}: ")
    set_font(r, size=12, bold=True, color=NAVY if kind != "caution" else "7A5A00")
    r = p.add_run(text)
    set_font(r, size=12, color=INK)


def add_heading(doc: Document, text: str, level=1, page_break=False) -> None:
    if page_break:
        doc.add_page_break()
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.first_line_indent = Cm(0)
    r = p.add_run(text)
    size = 16 if level == 1 else 14 if level == 2 else 13
    color = NAVY if level != 2 else BLUE
    set_font(r, size=size, bold=True, color=color)


def add_equation(doc: Document, equation: str, note: str | None = None) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(equation)
    set_font(r, size=12, italic=True, color=NAVY, name="Cambria Math")
    if note:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.space_after = Pt(7)
        r = p.add_run(note)
        set_font(r, size=10.5, italic=True, color=MUTED)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths_cm: list[float], caption: str | None = None, font_size=10.5) -> None:
    if caption:
        p = doc.add_paragraph(style="Table Caption")
        p.add_run(caption)
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths_cm)
    set_table_borders(table)
    set_repeat_table_header(table.rows[0])
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, NAVY)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(header)
        set_font(r, size=font_size, bold=True, color=WHITE)
    for ridx, values in enumerate(rows):
        cells = table.add_row().cells
        for idx, value in enumerate(values):
            cell = cells[idx]
            if ridx % 2 == 1:
                set_cell_shading(cell, PALE_BLUE)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if idx == 0 else WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(str(value))
            set_font(r, size=font_size, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_figure(doc: Document, path: Path, caption: str, width_cm=15.8) -> None:
    if not path.exists():
        add_callout(doc, "Thiếu hình", f"Không tìm thấy tệp {path.name} tại thời điểm tạo báo cáo.", "caution")
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run()
    shape = r.add_picture(str(path), width=Cm(width_cm))
    # Store meaningful alternative text in the DOCX drawing metadata so the
    # report remains understandable when images cannot be perceived.
    shape._inline.docPr.set("descr", caption)
    shape._inline.docPr.set("title", path.stem)
    p = doc.add_paragraph(style="Figure Caption")
    p.add_run(caption)


def pct(value: str | float, digits=2) -> str:
    return f"{float(value) * 100:.{digits}f}%"


def make_pipeline_diagram(path: Path) -> None:
    canvas = Image.new("RGB", (2376, 756), "white")
    draw = ImageDraw.Draw(canvas)
    font_dir = Path("C:/Windows/Fonts")
    bold = ImageFont.truetype(str(font_dir / "arialbd.ttf"), 27)
    body = ImageFont.truetype(str(font_dir / "arial.ttf"), 23)
    small = ImageFont.truetype(str(font_dir / "arial.ttf"), 21)
    stages = [
        (45, 170, 315, 190, "3CAD-ANI\nảnh + mask", "#EAF1F7"),
        (420, 170, 315, 190, "Audit & frozen split\nTrain / Val / Test", "#F4F7FA"),
        (795, 170, 360, 190, "Huấn luyện độc lập\nU-Net | SegFormer\nVMamba", "#EAF1F7"),
        (1215, 170, 315, 190, "Suy luận toàn ảnh\ntile 512\nstride 256", "#F4F7FA"),
        (1590, 170, 315, 190, "Hiệu chỉnh trên Val\nthreshold + policy", "#FFF4D6"),
        (1965, 170, 360, 190, "Test cố định\nPASS / DEFECT", "#E8F3EF"),
    ]
    for x, y, w, h, text, fill in stages:
        draw.rounded_rectangle((x, y, x + w, y + h), radius=18, fill=fill, outline="#2E5F8A", width=3)
        lines = text.split("\n")
        heights = [draw.textbbox((0, 0), line, font=bold)[3] for line in lines]
        y_text = y + (h - sum(heights) - 8 * (len(lines) - 1)) / 2
        for line, line_h in zip(lines, heights):
            box = draw.textbbox((0, 0), line, font=bold)
            draw.text((x + (w - (box[2] - box[0])) / 2, y_text), line, font=bold, fill="#1F2933")
            y_text += line_h + 8
    for idx in range(len(stages) - 1):
        x1 = stages[idx][0] + stages[idx][2]
        x2 = stages[idx + 1][0]
        y_mid = 265
        draw.line((x1 + 8, y_mid, x2 - 16, y_mid), fill="#5B6770", width=4)
        draw.polygon([(x2 - 16, y_mid - 10), (x2 - 16, y_mid + 10), (x2 - 2, y_mid)], fill="#5B6770")
    evidence = "Luồng bằng chứng: config, hash split, checkpoint, probability cache, metric, biểu đồ và test"
    box = draw.textbbox((0, 0), evidence, font=body)
    draw.text(((2376 - (box[2] - box[0])) / 2, 425), evidence, font=body, fill="#17365D")
    note = "Nguyên tắc chống rò rỉ dữ liệu: Test chỉ dùng để báo cáo, không điều chỉnh"
    note_box = draw.textbbox((0, 0), note, font=small)
    nx = (2376 - (note_box[2] - note_box[0])) / 2
    draw.rounded_rectangle((nx - 24, 520, nx + (note_box[2] - note_box[0]) + 24, 590), radius=14, fill="#FFF4D6", outline="#B8862B", width=3)
    draw.text((nx, 537), note, font=small, fill="#7A5A00")
    canvas.save(path)


def make_split_diagram(path: Path, scope: dict) -> None:
    splits = scope["protocol_preflight"]["splits"]
    names = ["Train", "Validation", "Test"]
    good = [splits["train"]["good"], splits["val"]["good"], splits["test"]["good"]]
    defect = [splits["train"]["defect"], splits["val"]["defect"], splits["test"]["defect"]]
    canvas = Image.new("RGB", (1692, 846), "white")
    draw = ImageDraw.Draw(canvas)
    font_dir = Path("C:/Windows/Fonts")
    title_font = ImageFont.truetype(str(font_dir / "arialbd.ttf"), 38)
    label_font = ImageFont.truetype(str(font_dir / "arial.ttf"), 28)
    value_font = ImageFont.truetype(str(font_dir / "arialbd.ttf"), 25)
    title = "Phân bố dữ liệu thực nghiệm 3CAD-ANI"
    title_box = draw.textbbox((0, 0), title, font=title_font)
    draw.text(((1692 - (title_box[2] - title_box[0])) / 2, 35), title, font=title_font, fill="#17365D")
    chart_top, chart_bottom = 130, 700
    max_total = max(g + d for g, d in zip(good, defect))
    for i in range(6):
        y = chart_bottom - i * (chart_bottom - chart_top) / 5
        draw.line((120, y, 1580, y), fill="#D9E1E8", width=2)
    centers = [390, 845, 1300]
    bar_w = 230
    for idx, (name, g, d) in enumerate(zip(names, good, defect)):
        total = g + d
        total_h = (total / max_total) * (chart_bottom - chart_top)
        good_h = (g / max_total) * (chart_bottom - chart_top)
        x0 = centers[idx] - bar_w / 2
        x1 = centers[idx] + bar_w / 2
        y_total = chart_bottom - total_h
        y_good = chart_bottom - good_h
        draw.rectangle((x0, y_good, x1, chart_bottom), fill="#3E73B8")
        draw.rectangle((x0, y_total, x1, y_good), fill="#2F8F76")
        for value, y_center in ((str(g), (y_good + chart_bottom) / 2), (str(d), (y_total + y_good) / 2)):
            box = draw.textbbox((0, 0), value, font=value_font)
            draw.text((centers[idx] - (box[2] - box[0]) / 2, y_center - 14), value, font=value_font, fill="white")
        total_text = str(total)
        box = draw.textbbox((0, 0), total_text, font=value_font)
        draw.text((centers[idx] - (box[2] - box[0]) / 2, y_total - 42), total_text, font=value_font, fill="#1F2933")
        box = draw.textbbox((0, 0), name, font=label_font)
        draw.text((centers[idx] - (box[2] - box[0]) / 2, 735), name, font=label_font, fill="#1F2933")
    draw.rectangle((1165, 91, 1195, 121), fill="#3E73B8")
    draw.text((1205, 90), "Good", font=label_font, fill="#1F2933")
    draw.rectangle((1350, 91, 1380, 121), fill="#2F8F76")
    draw.text((1390, 90), "Defect", font=label_font, fill="#1F2933")
    canvas.save(path)


def build_report() -> Path:
    WORK_DIR.mkdir(parents=True, exist_ok=True)
    REPORT_DIR.mkdir(parents=True, exist_ok=True)

    scope = load_json(THESIS_DIR / "scope.json")
    base_models = load_csv(THESIS_DIR / "tables" / "01_e2_architecture_comparison.csv")
    by_size = load_csv(THESIS_DIR / "tables" / "02_e3_defect_size.csv")
    by_group = load_csv(THESIS_DIR / "tables" / "03_e4_defect_group.csv")
    multi_region = load_csv(THESIS_DIR / "tables" / "04_e5_multi_region.csv")
    thresholds = load_csv(THESIS_DIR / "tables" / "05_e7_thresholds_from_validation.csv")
    spatial_uv = load_json(DECISION_DIR / "spatial" / "unet_vmamba" / "test" / "decision_metrics.json")
    spatial_uv_fnr_en = pct(spatial_uv["defect_fnr"])
    spatial_uv_fpr_en = pct(spatial_uv["defect_fpr"])
    spatial_uv_review_en = pct(spatial_uv["overall_review_rate"])
    spatial_uv_dice_en = pct(spatial_uv["positive_dice"])
    spatial_uv_fnr = spatial_uv_fnr_en.replace(".", ",")
    spatial_uv_fpr = spatial_uv_fpr_en.replace(".", ",")
    spatial_uv_review = spatial_uv_review_en.replace(".", ",")
    spatial_uv_dice = spatial_uv_dice_en.replace(".", ",")
    spatial_uv_fpr_reduction = f'{21.67 - float(spatial_uv["defect_fpr"]) * 100:.2f}'.replace(".", ",")
    spatial_uv_fnr_increase = f'{float(spatial_uv["defect_fnr"]) * 100 - 1.27:.2f}'.replace(".", ",")

    training_runs = [
        training_run("U-Net/ResNet18", Path("unet/TTTN/results/unet_r18/main_seed42"), "U-Net/ResNet18"),
        training_run("SegFormer-B0", Path("segformer/TTTN/results/segformer_b0/main_seed42"), "SegFormer-B0"),
        training_run("VMamba-T", Path("mamba/results/vmamba_t_s2l5/vmamba_b8_main_seed42"), "VMamba-T"),
    ]

    pipeline_png = WORK_DIR / "system_pipeline.png"
    split_png = WORK_DIR / "dataset_split.png"
    make_pipeline_diagram(pipeline_png)
    make_split_diagram(split_png, scope)

    doc = Document()
    configure_document(doc)
    set_running_furniture(doc.sections[0])
    add_cover(doc)

    doc.add_page_break()
    add_page_title(doc, "THÔNG TIN CẦN HOÀN THIỆN TRƯỚC KHI NỘP")
    add_callout(doc, "Lưu ý", "Các trường bôi vàng trên trang bìa và trong phần xác nhận phải được thay bằng thông tin thật. Báo cáo này được xây dựng theo một hệ trình bày học thuật chung vì workspace chưa có biểu mẫu DOCX chính thức của khoa.", "caution")
    add_bullets(doc, [
        "Thay tên trường, khoa/bộ môn, sinh viên, mã số sinh viên, lớp, đơn vị thực tập và người hướng dẫn.",
        "Đối chiếu quy định của khoa về lề trang, cỡ chữ, cách đánh số chương, biểu mẫu nhận xét và trang ký xác nhận.",
        "Điền thời gian thực tập, nhiệm vụ thực tế, mức độ đóng góp cá nhân và xác nhận của đơn vị; không sử dụng thông tin suy đoán.",
        "Kiểm tra quyền sử dụng hoặc quyền tái phân phối dữ liệu trước khi đính kèm ảnh gốc 3CAD vào bản công khai.",
    ])

    doc.add_page_break()
    add_page_title(doc, "LỜI CAM ĐOAN")
    add_para(doc, "Tôi cam đoan báo cáo này phản ánh trung thực quá trình thực tập và phần công việc do tôi thực hiện. Các kết quả định lượng được trích từ artifact đã đóng băng của dự án; các nguồn học thuật và phần mềm được ghi nhận trong danh mục tài liệu tham khảo. Những nội dung sử dụng công cụ trí tuệ nhân tạo phải được khai báo theo quy định của cơ sở đào tạo và được người viết kiểm tra lại bằng mã nguồn, dữ liệu, log và kiểm thử.")
    add_para(doc, "Tôi chịu trách nhiệm về tính chính xác của thông tin cá nhân, thời gian thực tập, sự phân công tại đơn vị và phạm vi đóng góp được bổ sung vào bản cuối.")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Pt(36)
    r = p.add_run("TP. Hồ Chí Minh, ngày ...... tháng ...... năm 2026\nSinh viên thực hiện\n\n\n[ĐIỀN HỌ VÀ TÊN]")
    set_font(r, size=13, color=INK)
    r.font.highlight_color = WD_COLOR_INDEX.YELLOW

    doc.add_page_break()
    add_page_title(doc, "LỜI CẢM ƠN")
    add_para(doc, "Báo cáo trân trọng ghi nhận sự hướng dẫn của giảng viên, sự hỗ trợ của người hướng dẫn tại đơn vị thực tập và các góp ý chuyên môn trong quá trình xây dựng hệ thống. Người viết cần bổ sung tên cá nhân, đơn vị và đóng góp thực tế trước khi nộp để phần cảm ơn phản ánh đúng quá trình làm việc.")
    add_para(doc, "Dự án sử dụng tập dữ liệu 3CAD do Yang và cộng sự công bố, cùng các công trình nền tảng U-Net, ResNet, SegFormer và VMamba. Tác giả cảm ơn cộng đồng mã nguồn mở đã cung cấp các thư viện phục vụ huấn luyện, đánh giá và triển khai.")

    doc.add_page_break()
    add_page_title(doc, "TÓM TẮT")
    add_para(doc, "Kiểm tra khuyết tật bề mặt trong sản xuất linh kiện nhôm đòi hỏi đồng thời khả năng phát hiện bất thường nhỏ và kiểm soát báo động giả. Báo cáo trình bày quá trình xây dựng một hệ thống phân đoạn nhị phân trên tập con Aluminum New Ipad (ANI) của 3CAD, so sánh ba họ kiến trúc U-Net/ResNet18, SegFormer-B0 và VMamba-T, đồng thời phát triển lớp hỗ trợ quyết định nhằm chuyển bản đồ xác suất thành kết luận PASS hoặc DEFECT.")
    add_para(doc, f"Dữ liệu được kiểm tra tính toàn vẹn và cố định thành 5.733 ảnh Train, 718 ảnh Validation và 717 ảnh Test. Ba mô hình được huấn luyện độc lập trên patch 512 x 512; Validation và Test dùng suy luận toàn ảnh theo cửa sổ trượt với stride 256. Checkpoint và ngưỡng được chọn chỉ trên Validation, sau đó khóa trước khi đánh giá Test. Kết quả cho thấy VMamba-T đạt Positive Dice 76,16%, IoU 64,78%, image F1 91,21% và FPR 21,67%, tốt nhất về chất lượng mặt nạ trong ba mô hình. Cấu hình Spatial ensemble U-Net + VMamba đạt FNR {spatial_uv_fnr}, FPR {spatial_uv_fpr}, Positive Dice {spatial_uv_dice} và review rate {spatial_uv_review} trên Test. So với ngưỡng pixel gốc, policy giảm báo động giả nhưng chuyển một phần ca bất định sang REVIEW.")
    add_para(doc, "Hệ thống hoàn chỉnh gồm mã lõi tái sử dụng, quy trình train/evaluation có bằng chứng, công cụ rà soát dữ liệu local-first, API FastAPI và giao diện web React. Kết quả kiểm thử hiện tại gồm protocol preflight PASS, 14 kiểm thử Python PASS, production web build PASS và 2 kiểm thử Node PASS. Báo cáo đồng thời nêu rõ giới hạn về cỡ mẫu của một số nhóm lỗi, sự không đồng nhất cấu hình huấn luyện cuối, phạm vi thí nghiệm chưa thực hiện và yêu cầu xác minh quyền phân phối dữ liệu.")
    add_para(doc, "Từ khóa: phân đoạn ảnh; khuyết tật bề mặt nhôm; 3CAD; U-Net; SegFormer; VMamba; hỗ trợ quyết định; kiểm soát báo động giả.", bold_lead="Từ khóa:", indent=False)

    doc.add_page_break()
    add_page_title(doc, "ABSTRACT")
    add_para(doc, "Industrial aluminum-surface inspection requires both sensitivity to small defects and control of false alarms. This report presents the development of a binary segmentation system on the Aluminum New Ipad subset of 3CAD. Three architectures - U-Net/ResNet18, SegFormer-B0, and VMamba-T - are evaluated under a frozen Train/Validation/Test protocol. A decision layer converts pixel-level probability maps into operational PASS or DEFECT outcomes.")
    add_para(doc, f"The experimental dataset contains 5,733 training images, 718 validation images, and 717 test images. Models are trained on 512 x 512 patches, while full-resolution validation and testing use sliding-window inference with a stride of 256. Model checkpoints, segmentation thresholds, and decision policies are selected exclusively on Validation and frozen before Test evaluation. VMamba-T obtains the strongest mask quality, with 76.16% Positive Dice and 64.78% Positive IoU. The selected U-Net + VMamba spatial ensemble reaches {spatial_uv_fnr_en} false-negative rate, {spatial_uv_fpr_en} false-positive rate, {spatial_uv_dice_en} Positive Dice, and {spatial_uv_review_en} review rate on Test.")
    add_para(doc, "The delivered system also includes reproducibility evidence, a local-first dataset review tool, a FastAPI inference service, and a React web interface. The report explicitly documents limitations, unresolved data-licensing questions, and the distinction between validated results and pending external verification.")
    add_para(doc, "Keywords: semantic segmentation; aluminum surface defect; 3CAD; U-Net; SegFormer; VMamba; decision support; false-alarm control.", bold_lead="Keywords:", indent=False)

    doc.add_page_break()
    add_page_title(doc, "DANH MỤC TỪ VIẾT TẮT")
    add_table(doc, ["Từ viết tắt", "Diễn giải"], [
        ["ANI", "Aluminum New Ipad - nhóm sản phẩm nhôm được sử dụng trong 3CAD"],
        ["AI", "Artificial Intelligence - trí tuệ nhân tạo"],
        ["API", "Application Programming Interface"],
        ["AUPRC", "Area Under the Precision-Recall Curve"],
        ["AUROC", "Area Under the Receiver Operating Characteristic Curve"],
        ["BCE", "Binary Cross-Entropy"],
        ["CNN", "Convolutional Neural Network"],
        ["FNR / FPR", "False Negative Rate / False Positive Rate"],
        ["FPN", "Feature Pyramid Network"],
        ["GT", "Ground Truth"],
        ["IoU", "Intersection over Union"],
        ["OOF", "Out-of-fold - dự báo ngoài fold trong cross-validation"],
        ["ROI", "Region of Interest"],
        ["SS2D", "2D Selective Scan trong VMamba"],
    ], [3.2, 12.8], font_size=10.5)

    doc.add_page_break()
    add_page_title(doc, "MỤC LỤC")
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    add_field(p, 'TOC \\o "1-3" \\h \\z \\u', "Nhấn F9 trong Microsoft Word để cập nhật mục lục nếu cần.")
    add_callout(doc, "Hướng dẫn", "Sau khi thay đổi nội dung hoặc thông tin cá nhân, chọn toàn bộ tài liệu (Ctrl+A) và nhấn F9 để cập nhật mục lục và số trang.")

    add_heading(doc, "CHƯƠNG 1. TỔNG QUAN ĐỀ TÀI", level=1, page_break=True)
    add_heading(doc, "1.1. Bối cảnh và lý do chọn đề tài", level=2)
    add_para(doc, "Trong dây chuyền sản xuất sản phẩm điện tử và linh kiện kim loại, các khuyết tật như va đập, bầm, xước, vết dao hoặc nhiều vùng lỗi đồng thời có thể ảnh hưởng đến thẩm mỹ và chất lượng sản phẩm. Kiểm tra thủ công phụ thuộc vào kinh nghiệm người vận hành, khó duy trì ổn định khi số lượng ảnh lớn và đặc biệt dễ bỏ sót những vùng bất thường nhỏ, tương phản thấp.")
    add_para(doc, "Học sâu cho phép học trực tiếp đặc trưng từ ảnh và tạo mặt nạ ở mức pixel. Tuy nhiên, một mô hình phân đoạn tốt chưa tự động đồng nghĩa với một hệ thống kiểm tra tốt: threshold thấp giúp giảm bỏ sót nhưng thường làm tăng báo động giả; threshold cao tạo mặt nạ sạch hơn nhưng có nguy cơ bỏ qua lỗi nhỏ. Vì vậy, đề tài không chỉ so sánh mô hình mà còn tổ chức dữ liệu, khóa quy trình thí nghiệm, kiểm soát rò rỉ dữ liệu và xây dựng lớp quyết định phù hợp với ngữ cảnh vận hành.")
    add_callout(doc, "Tên đề tài đề xuất", "Xây dựng hệ thống phân đoạn và hỗ trợ quyết định phát hiện khuyết tật bề mặt nhôm dựa trên học sâu.")

    add_heading(doc, "1.2. Mục tiêu nghiên cứu", level=2)
    add_bullets(doc, [
        "Xây dựng pipeline phân đoạn khuyết tật bề mặt nhôm có thể tái tạo, từ dữ liệu đến huấn luyện, đánh giá và triển khai.",
        "So sánh ba đại diện kiến trúc: CNN encoder-decoder (U-Net/ResNet18), Transformer phân cấp (SegFormer-B0) và state-space model thị giác (VMamba-T).",
        "Đánh giá ở mức ảnh, pixel và vùng connected component; phân tích theo kích thước, nhóm lỗi và số vùng lỗi.",
        "Hiệu chỉnh threshold và policy chỉ trên Validation, giữ Test làm tập báo cáo độc lập.",
        "Giảm báo động giả bằng hậu xử lý không gian và bộ xác minh học được; tích hợp cấu hình cuối vào API và giao diện web.",
        "Cung cấp công cụ rà soát dữ liệu, log và bằng chứng kiểm thử phục vụ bàn giao thực tập.",
    ])

    add_heading(doc, "1.3. Câu hỏi nghiên cứu", level=2)
    add_bullets(doc, [
        "RQ1: Kiến trúc nào tạo mặt nạ khuyết tật tốt nhất trên tập Test cố định của 3CAD-ANI?",
        "RQ2: Hiệu năng thay đổi như thế nào theo kích thước, nhóm khuyết tật và số lượng vùng lỗi?",
        "RQ3: Có thể giảm báo động giả ở mức ảnh mà vẫn duy trì tỷ lệ bỏ sót trong giới hạn chấp nhận được hay không?",
        "RQ4: Hệ thống có đủ bằng chứng để người khác kiểm tra cấu hình, split, checkpoint, metric và luồng triển khai hay không?",
    ], numbered=True)

    add_heading(doc, "1.4. Đối tượng, phạm vi và giới hạn", level=2)
    add_para(doc, "Đối tượng nghiên cứu là bài toán phân đoạn nhị phân trên ảnh bề mặt nhôm thuộc nhóm ANI. Mặt nạ đầu ra chỉ phân biệt nền và bất thường; tên nhóm lỗi được dùng làm metadata để phân tích, không phải nhãn đa lớp. Scope kết quả cuối gồm E0, E2, E3, E4, E5, E7 và E8. Các thí nghiệm E1, E6, E9, E10, E11 không thuộc phạm vi bàn giao hiện tại.")
    add_callout(doc, "Giới hạn diễn giải", "Số liệu độ trễ và VRAM chỉ mang tính mô tả, không được xem là so sánh E6 chính thức vì cấu hình batch, số epoch và điều kiện đo của các run cuối chưa hoàn toàn đồng nhất.", "caution")

    add_heading(doc, "1.5. Đóng góp chính của dự án", level=2)
    add_bullets(doc, [
        "Repository được tách rõ src, scripts, apps, tests, data, artifacts, runtime và submission; các đường dẫn launcher ở gốc là giao diện ổn định.",
        "Dữ liệu 3CAD-ANI được audit, loại một bản ghi trùng/thiếu và đóng băng split bằng SHA-256.",
        "Ba mô hình được đánh giá theo cùng định nghĩa metric và cùng tập Test.",
        "Decision pipeline kết hợp ROI biên tối, connected components, luật Adaptive và Spatial ensemble.",
        "Hai ứng dụng hỗ trợ sử dụng thực tế: Dataset Review Studio và web demo dự báo.",
        "Bộ kiểm thử tự động xác minh protocol, logic quyết định, trạng thái huấn luyện, review tool và giao diện web.",
    ])

    add_heading(doc, "CHƯƠNG 2. CƠ SỞ LÝ THUYẾT VÀ CÔNG TRÌNH LIÊN QUAN", level=1, page_break=True)
    add_heading(doc, "2.1. Bài toán phân đoạn khuyết tật", level=2)
    add_para(doc, "Với ảnh đầu vào x có kích thước H x W x 3, mô hình f_theta tạo logit z = f_theta(x). Sau phép sigmoid, mỗi pixel nhận xác suất p thuộc [0,1]. Mặt nạ nhị phân được tạo bởi quy tắc p >= t, trong đó t là threshold được chọn trên Validation. Ở mức ảnh, dự án dùng max(probability map) làm điểm bất thường cơ sở; decision policy nâng quy tắc này lên bằng thông tin vùng, vị trí và đồng thuận giữa mô hình.")
    add_equation(doc, "p(i,j) = sigmoid(z(i,j));   y_hat(i,j) = 1 nếu p(i,j) >= t, ngược lại bằng 0.")

    add_heading(doc, "2.2. U-Net và ResNet18", level=2)
    add_para(doc, "U-Net sử dụng đường co để thu nhận ngữ cảnh và đường giãn đối xứng để khôi phục định vị; skip connection giúp truyền đặc trưng không gian từ encoder sang decoder [2]. Trong dự án, encoder được thay bằng ResNet18 pretrained ImageNet. Residual connection của ResNet giúp tối ưu mạng sâu thông qua học phần dư [3]. Đây là baseline CNN có cấu trúc dễ giải thích và tốc độ suy luận cao.")

    add_heading(doc, "2.3. SegFormer-B0", level=2)
    add_para(doc, "SegFormer kết hợp Transformer encoder phân cấp với decoder MLP nhẹ; thiết kế không dùng positional encoding tuyệt đối và hợp nhất đặc trưng đa tỉ lệ [4]. Bản B0 được chọn để giữ kích thước mô hình vừa phải. Đầu phân đoạn nhị phân được khởi tạo mới, trong khi encoder sử dụng pretraining ImageNet.")

    add_heading(doc, "2.4. VMamba-T", level=2)
    add_para(doc, "VMamba đưa state-space model vào thị giác bằng các khối Visual State-Space và mô-đun SS2D quét theo nhiều hướng, nhằm thu ngữ cảnh toàn cục với độ phức tạp tuyến tính theo chuỗi đầu vào [5]. Dự án dùng VMamba-T s2l5 làm backbone và một decoder FPN nhẹ cho phân đoạn nhị phân. Mô hình này đạt chất lượng mặt nạ cao nhất nhưng chậm hơn đáng kể trong runtime hiện tại.")

    add_heading(doc, "2.5. Hàm mất mát và chiến lược tối ưu", level=2)
    add_para(doc, "Loss huấn luyện là tổ hợp 0,5 BCE trên mọi mẫu và 0,5 Dice loss chỉ trên mẫu có ground-truth dương. BCE cho phép ảnh Good tham gia học giảm false positive; Dice chỉ áp dụng cho mẫu dương để tránh mặt nạ toàn 0 trở thành mục tiêu chồng lấn không có ý nghĩa.")
    add_equation(doc, "L = 0,5 x BCE(y, p) + 0,5 x (1 - Dice(y, p)) trên mẫu dương.")
    add_equation(doc, "Dice = (2 x |Y giao P| + eps) / (|Y| + |P| + eps).")
    add_para(doc, "Bộ tối ưu AdamW tách weight decay khỏi cập nhật gradient thích nghi [6]. Dự án sử dụng learning rate encoder 1e-5, decoder 1e-4, weight decay 1e-4, gradient clipping 1,0, mixed precision trên CUDA và ReduceLROnPlateau theo Positive Dice@0.5 trên Validation.")

    add_heading(doc, "2.6. Thước đo đánh giá", level=2)
    add_table(doc, ["Nhóm", "Thước đo", "Ý nghĩa trong đề tài"], [
        ["Mức ảnh", "AUROC, AUPRC", "Khả năng xếp hạng Good/Defect trên nhiều threshold"],
        ["Mức ảnh", "Precision, Recall, Specificity, F1", "Chất lượng quyết định tại threshold đã khóa"],
        ["Sai số", "FNR, FPR", "Bỏ sót khuyết tật và báo động giả"],
        ["Mức pixel", "Positive Dice, Positive IoU", "Độ chồng lấn mặt nạ trên ảnh Defect"],
        ["Mức pixel", "Pixel AUPRC/Recall/Precision", "Chất lượng dự báo pixel toàn cục"],
        ["Mức vùng", "Region Recall", "Tỷ lệ connected component GT có giao với dự báo"],
        ["Vận hành", "Accuracy, review rate", "Chất lượng và mức tự động hóa của decision policy"],
    ], [2.4, 4.1, 9.5], caption="Bảng 2.1. Hệ thước đo được sử dụng", font_size=10)
    add_equation(doc, "FNR = FN / (TP + FN);   FPR = FP / (TN + FP);   IoU = |Y giao P| / |Y hop P|.")
    add_callout(doc, "Nguyên tắc", "Không dùng một metric duy nhất để kết luận. Với kiểm tra chất lượng, cần đọc đồng thời Dice/IoU, Recall/FNR và FPR vì các đại lượng này phản ánh các rủi ro khác nhau.")

    add_heading(doc, "CHƯƠNG 3. PHÂN TÍCH VÀ THIẾT KẾ HỆ THỐNG", level=1, page_break=True)
    add_heading(doc, "3.1. Yêu cầu chức năng và phi chức năng", level=2)
    add_table(doc, ["Mã", "Yêu cầu", "Tiêu chí chấp nhận"], [
        ["FR-01", "Nạp ảnh sản phẩm và sinh probability map", "Không trả dự báo giả khi checkpoint thiếu"],
        ["FR-02", "Xuất mặt nạ/overlay cho từng mô hình", "Kích thước khớp ảnh đầu vào"],
        ["FR-03", "Sinh kết luận PASS/DEFECT", "Dùng policy đã khóa, trả kèm lý do và đồng thuận"],
        ["FR-04", "Rà soát và sửa nhãn/mask", "Không ghi đè dữ liệu nguồn; có audit log"],
        ["FR-05", "Train và đánh giá tái lập", "Lưu config, snapshot split, hash, checkpoint, history"],
        ["NFR-01", "Chống rò rỉ dữ liệu", "Mọi lựa chọn dùng Validation; Test chỉ báo cáo"],
        ["NFR-02", "Khả năng kiểm tra", "Có test tự động và verification launcher"],
        ["NFR-03", "Khả năng bàn giao", "Đường dẫn tương đối, README và launcher ổn định"],
    ], [1.8, 6.0, 8.2], caption="Bảng 3.1. Yêu cầu cốt lõi của hệ thống", font_size=9.8)

    add_heading(doc, "3.2. Kiến trúc tổng thể", level=2)
    add_figure(doc, pipeline_png, "Hình 3.1. Luồng dữ liệu, huấn luyện, hiệu chỉnh và triển khai", width_cm=16.0)
    add_para(doc, "Luồng được thiết kế theo nguyên tắc bằng chứng đi cùng kết quả. Mỗi run huấn luyện lưu cấu hình, thông tin môi trường, mô tả mô hình, snapshot split và SHA-256, lịch sử epoch, checkpoint tốt nhất và cuối cùng, đường học và ảnh theo dõi Validation. Probability cache tách chi phí suy luận khỏi các thí nghiệm policy, giúp chạy lại logic quyết định mà không cần chạy model nhiều lần.")

    add_heading(doc, "3.3. Thiết kế dữ liệu", level=2)
    add_para(doc, "Nguồn dữ liệu là 3CAD, một bộ dữ liệu công nghiệp thực tế gồm 27.039 ảnh độ phân giải cao, tám nhóm sản phẩm và nhãn bất thường mức pixel [1]. Dự án sử dụng nhóm Aluminum New Ipad. Bản nguồn cục bộ được audit có 7.169 ảnh; một ảnh bị loại do file nguồn thiếu và trùng SHA-256 với ảnh liền kề, vì vậy tập thực nghiệm còn 7.168 ảnh.")
    add_figure(doc, split_png, "Hình 3.2. Phân bố Good/Defect theo frozen split", width_cm=14.8)
    splits = scope["protocol_preflight"]["splits"]
    add_table(doc, ["Split", "Tổng", "Good", "Defect", "Vai trò"], [
        ["Train", f'{splits["train"]["n"]:,}'.replace(",", "."), f'{splits["train"]["good"]:,}'.replace(",", "."), f'{splits["train"]["defect"]:,}'.replace(",", "."), "Học tham số và suy ra ngưỡng nhóm từ dữ liệu Train"],
        ["Validation", str(splits["val"]["n"]), str(splits["val"]["good"]), str(splits["val"]["defect"]), "Chọn checkpoint, threshold và policy"],
        ["Test", str(splits["test"]["n"]), str(splits["test"]["good"]), str(splits["test"]["defect"]), "Đánh giá cuối, không tinh chỉnh"],
    ], [2.5, 2.0, 2.0, 2.0, 7.5], caption="Bảng 3.2. Quy mô và vai trò của các split", font_size=10)
    add_callout(doc, "Bằng chứng toàn vẹn", "Kiểm tra overlap theo đường dẫn và nội dung đều PASS. SHA-256 của train.csv, val.csv và test.csv lần lượt bắt đầu bằng 9ba4ba9c, b33791cc và cf0372bd; giá trị đầy đủ nằm trong hồ sơ xác minh của dự án.")

    add_heading(doc, "3.4. Tiền xử lý và lấy mẫu", level=2)
    add_bullets(doc, [
        "Ảnh được đọc RGB và chuẩn hóa theo thống kê ImageNet.",
        "Train lấy patch 512 x 512 trực tiếp từ ảnh gốc: ảnh Defect dùng crop dương có nhận biết GT, ảnh Good dùng crop âm ngẫu nhiên.",
        "Tăng cường mặc định chỉ gồm biến đổi quang học nhẹ: brightness, contrast và gamma khoảng +/-10%.",
        "Validation/Test không dùng GT để chọn crop; suy luận toàn ảnh bằng tile 512, stride 256 và lấy trung bình ở vùng chồng lấn.",
        "Ảnh nhỏ hơn tile được pad và sau suy luận được cắt về kích thước gốc.",
    ])

    add_heading(doc, "3.5. Thiết kế decision policy", level=2)
    add_para(doc, "Decision policy nhận ảnh gốc và probability map của từng model. Đầu tiên, vùng gần như đen nối với biên ảnh được xem là padding/gá ngoài sản phẩm và loại khỏi ROI. Sau đó mask được tách thành connected components; vùng đủ lớn được xem là strong component, vùng nhỏ có xác suất rất cao được giữ làm tín hiệu REVIEW hoặc ứng viên cứu hộ. Khi nhiều model cùng phát hiện vùng gần nhau, spatial consensus tăng độ tin cậy của kết luận DEFECT.")
    add_bullets(doc, [
        "Spatial rule: kết hợp area, peak probability, vị trí và số model đồng thuận.",
        "Adaptive single-model rule: điều chỉnh yêu cầu component theo bằng chứng của từng ảnh.",
        "Spatial ensemble U-Net + VMamba: giữ kết luận đồng thuận theo vị trí của hai model; Test không tham gia chọn threshold.",
    ])

    add_heading(doc, "3.6. Thiết kế ứng dụng", level=2)
    add_para(doc, "Dataset Review Studio là công cụ local-first dùng FastAPI, JavaScript và SQLite. Dataset nguồn được mở chỉ đọc; mọi quyết định review, mask sửa và export được lưu riêng. Người dùng có thể lọc ảnh Good score cao, false positive, zero overlap hoặc model bất đồng; xem ảnh gốc, GT, overlay/heatmap/binary; sau đó duyệt, đổi nhãn, sửa mask, đánh dấu uncertain hoặc loại mẫu.")
    add_para(doc, "Web demo gồm backend FastAPI và frontend React. Endpoint /health báo tình trạng checkpoint, runtime và policy; endpoint /infer nhận ảnh, chạy các model khả dụng, áp dụng policy và trả quyết định, lý do, mask và overlay. Frontend không hiển thị model như khả dụng nếu checkpoint hoặc runtime thiếu.")

    add_heading(doc, "CHƯƠNG 4. CÀI ĐẶT VÀ HIỆN THỰC", level=1, page_break=True)
    add_heading(doc, "4.1. Tổ chức mã nguồn", level=2)
    add_table(doc, ["Thư mục", "Trách nhiệm"], [
        ["src/threecad_segmentation", "Kiến trúc model, loss, train loop, full-resolution evaluation và decision modules"],
        ["scripts/training", "Entry point huấn luyện U-Net, SegFormer, VMamba và retraining"],
        ["scripts/evaluation", "Đánh giá từng model trên Validation/Test"],
        ["scripts/experiments", "Probability cache, decision policy và batch thí nghiệm"],
        ["scripts/reporting", "Tổng hợp bảng, audit test case và trực quan hóa"],
        ["apps/dataset_review", "Rà soát/sửa dữ liệu không phá hủy"],
        ["apps/web_demo", "API suy luận và giao diện trình diễn"],
        ["tests", "Unit/smoke tests cho logic ML và review tool"],
        ["submission", "Báo cáo, log, checklist và bằng chứng"],
    ], [4.3, 11.7], caption="Bảng 4.1. Phân rã module của repository", font_size=10)

    add_heading(doc, "4.2. Hiện thực ba mô hình", level=2)
    add_para(doc, "UNetResNet18 tái sử dụng các tầng convolution và residual block của ResNet18 làm encoder; decoder gồm các block upsample, ghép skip feature và DoubleConv. SegFormerB0Binary lấy encoder nvidia/mit-b0 và thay decode head cho một lớp logit. VMambaTBinary bao bọc backbone VMamba-T và decoder FPN nhẹ, đưa nhiều mức feature về cùng không gian trước khi dự báo mask.")
    add_table(doc, ["Mô hình", "Backbone/encoder", "Decoder", "Pretraining"], [
        ["U-Net/ResNet18", "ResNet18", "U-Net decoder + skip connection", "ImageNet"],
        ["SegFormer-B0", "MiT-B0 Transformer phân cấp", "All-MLP binary head", "ImageNet"],
        ["VMamba-T", "VMamba-T s2l5 / SS2D", "Lightweight FPN binary head", "ImageNet"],
    ], [3.5, 4.0, 4.9, 3.6], caption="Bảng 4.2. Kiến trúc được hiện thực", font_size=9.8)

    add_heading(doc, "4.3. Vòng lặp huấn luyện và quản lý trạng thái", level=2)
    add_para(doc, "TrainConfig tập trung các tham số dữ liệu, tối ưu và bằng chứng đầu ra. Seed được đặt cho Python, NumPy, PyTorch và worker. Encoder và decoder nhận learning rate riêng. Gradient accumulation cho phép điều chỉnh effective batch theo giới hạn VRAM; AMP và gradient clipping được bật khi phù hợp. Sau mỗi epoch, hệ thống tính Validation toàn ảnh, cập nhật scheduler, lưu best/last checkpoint và giám sát early stopping.")
    add_callout(doc, "Tính tái tạo", "Mỗi run lưu config.json, environment.json, model_info.json, split snapshot, SHA-256, training_history.csv, training_summary.json, checkpoint và các đường học. Khi bị gián đoạn, interrupt checkpoint giữ trạng thái để tiếp tục.")

    add_heading(doc, "4.4. Đánh giá toàn ảnh", level=2)
    add_para(doc, "Hàm predict_sliding_window tạo danh sách vị trí tile bảo đảm phủ biên ảnh; dự báo từng batch tile; cộng probability và đếm số lần phủ; sau đó chia trung bình tại mỗi pixel. evaluate_split vừa tính metric ảnh, pixel và vùng, vừa lưu confusion matrix, ROC/PR curve, threshold scan, định tính và thời gian. Các bin kích thước và số component được suy ra từ Train để tránh dùng Test thiết kế nhóm.")

    add_heading(doc, "4.5. Kiểm soát dữ liệu và quyền riêng tư", level=2)
    add_para(doc, "Dataset Review Studio dùng SQLite để lưu trạng thái và lịch sử. Mask sửa được lưu dưới edits/masks; export chứa manifest sạch, split, corrected masks, audit log, review_events, hard negatives và unresolved. Mẫu uncertain, exclude, mask Defect rỗng hoặc Defect mới chưa có mask không bị đưa âm thầm vào bản sạch.")
    add_para(doc, "Git chỉ theo dõi mã nguồn, cấu hình, tài liệu, notebook và hồ sơ nộp. Dataset, checkpoint, probability cache, môi trường ảo, runtime wheel và archive được loại khỏi Git. Cách tổ chức này giảm nguy cơ commit nhầm dữ liệu lớn nhưng không thay thế việc xác minh license trước khi chia sẻ.")

    add_heading(doc, "4.6. Công nghệ sử dụng", level=2)
    add_table(doc, ["Lớp", "Công nghệ", "Vai trò"], [
        ["ML", "Python, PyTorch, Transformers, timm", "Model, train, inference"],
        ["Xử lý dữ liệu", "NumPy, Pandas, OpenCV, Pillow", "Ảnh, mask, thống kê và CSV"],
        ["Backend", "FastAPI, Uvicorn", "API review và inference"],
        ["Frontend", "React, TypeScript, Vite/vinext", "Giao diện web"],
        ["Kiểm thử", "pytest/unittest, Node test", "ML, review tool và rendered HTML"],
        ["Môi trường train", "Kaggle GPU, CUDA, AMP", "Huấn luyện và đánh giá mô hình"],
    ], [3.1, 5.0, 7.9], caption="Bảng 4.3. Ngăn xếp công nghệ", font_size=10)

    add_heading(doc, "CHƯƠNG 5. THỰC NGHIỆM VÀ KẾT QUẢ", level=1, page_break=True)
    add_heading(doc, "5.1. Giao thức thí nghiệm", level=2)
    add_bullets(doc, [
        "Sử dụng train.csv, val.csv, test.csv cố định; kiểm tra overlap, nội dung và file trước khi chạy.",
        "Chọn checkpoint tốt nhất theo Validation Positive Dice@0.5.",
        "Quét threshold 0,05 đến 0,95 với bước 0,01 trên Validation; giữ ứng viên có FNR <= 10%, tối đa hóa Positive Dice và dùng FPR thấp hơn để phá hòa.",
        "Khóa threshold/policy; áp dụng đúng một lần lên Test.",
        "Phân tích E3/E5 dùng ngưỡng nhóm suy ra từ connected components của Train.",
        "Các nhóm dưới 10 ảnh Test được đánh dấu mô tả; knife mark chỉ có 2 ảnh.",
    ])
    add_table(doc, ["Mô hình", "Best epoch", "Val Dice@0.5", "Threshold", "Batch / Accum", "Max epoch"], [
        ["U-Net/ResNet18", "48", "70,22%", "0,49", "38 / 1", "50"],
        ["SegFormer-B0", "50", "69,20%", "0,66", "16 / 2", "50"],
        ["VMamba-T", "23", "76,84%", "0,51", "16 / 1", "25"],
    ], [3.3, 2.1, 2.7, 2.0, 3.1, 2.8], caption="Bảng 5.1. Artifact mô hình cuối và ngưỡng chọn trên Validation", font_size=9.5)
    add_callout(doc, "Sai khác cần công khai", "Batch/effective batch và số epoch của artifact cuối khác nhau giữa mô hình. Vì vậy so sánh chính tập trung vào chất lượng dự báo; không kết luận kiến trúc nào huấn luyện hiệu quả hơn.", "caution")

    add_heading(doc, "5.1.1. Hồ sơ huấn luyện và đường cong học", level=3)
    add_para(doc, "Sau khi tải lại output từ Kaggle, báo cáo có thể đối chiếu trực tiếp lịch sử theo từng epoch thay vì chỉ ghi nhận best checkpoint. Bảng dưới đây lấy từ training_summary.json và training_history.csv của đúng các run đã dùng cho kết quả cuối; các đường cong minh họa đồng thời train/validation loss và Positive Dice.")
    training_rows = []
    for run in training_runs:
        summary = run["summary"]
        training_rows.append([
            run["model"],
            str(summary.get("epochs_completed", "-")),
            str(summary.get("best_epoch", "-")),
            pct(summary["best_val_positive_dice_0.5"], 2) if "best_val_positive_dice_0.5" in summary else "-",
            str(summary.get("effective_batch_size", "-")),
            f'{float(summary["max_peak_vram_gb"]):.1f} GB' if "max_peak_vram_gb" in summary else "-",
            f'{float(summary["total_train_run_wall_seconds"]) / 3600:.1f} h' if "total_train_run_wall_seconds" in summary else "-",
        ])
    add_table(doc, ["Mô hình", "Epoch hoàn tất", "Best epoch", "Best Val Dice", "Effective batch", "Peak VRAM", "Wall time"], training_rows, [3.0, 2.0, 1.9, 2.3, 2.5, 2.1, 2.2], caption="Bảng 5.2. Tóm tắt quá trình huấn luyện từ output Kaggle", font_size=8.8)
    for index, run in enumerate(training_runs, start=1):
        add_figure(doc, run["curve"], f'Hình 5.1.{index}. Đường cong huấn luyện {run["label"]}', width_cm=15.4)
    add_para(doc, "Ngoài các hình trên, mỗi run còn lưu loss_components.png, learning_rate.png, epoch_time.png, vram.png, environment.json và best.pt. Toàn bộ hồ sơ được giữ tại artifacts/training_download để kiểm tra lại cấu hình, môi trường và trạng thái huấn luyện mà không làm thay đổi các metric Test đã khóa.")

    add_heading(doc, "5.2. So sánh tổng thể ba kiến trúc (E2)", level=2)
    add_figure(doc, FIG_DIR / "01_e2_overall_metrics.png", "Hình 5.1. So sánh hiệu năng tổng thể trên tập Test", width_cm=16.0)
    results_rows = []
    for row in base_models:
        results_rows.append([
            row["model"], pct(row["image_auroc"], 1), pct(row["image_auprc"], 1),
            pct(row["image_f1"], 1), pct(row["positive_dice"], 1), pct(row["positive_iou"], 1), pct(row["region_recall_any_overlap"], 1)
        ])
    add_table(doc, ["Mô hình", "AUROC", "AUPRC", "F1 ảnh", "Dice", "IoU", "Recall vùng"], results_rows, [3.4, 2.0, 2.0, 2.0, 2.1, 2.1, 2.4], caption="Bảng 5.2. Kết quả chính trên 717 ảnh Test", font_size=9.1)
    add_para(doc, "VMamba-T dẫn đầu AUROC (94,58%), image F1 (91,21%), Positive Dice (76,16%) và Positive IoU (64,78%). SegFormer-B0 đạt recall ảnh cao nhất 99,75% và chỉ bỏ sót một ảnh, nhưng FPR 61,92% làm F1 giảm còn 79,64%. U-Net cân bằng hơn SegFormer nhưng chất lượng mặt nạ thấp hơn VMamba.")

    add_figure(doc, FIG_DIR / "04_e2_pixel_metrics.png", "Hình 5.2. Chất lượng dự báo ở mức pixel", width_cm=15.7)
    add_para(doc, "Ở mức pixel, VMamba-T tiếp tục dẫn đầu với Pixel AUPRC 79,13%, Recall 79,74% và Precision 80,90%. Điều này phù hợp với Positive Dice/IoU cao hơn, cho thấy lợi thế không chỉ đến từ threshold mức ảnh mà còn từ probability map có chất lượng tốt hơn.")

    add_heading(doc, "5.3. Phân tích nhầm lẫn và chi phí sai số", level=2)
    confusion_rows = []
    for row in base_models:
        confusion_rows.append([
            row["model"], row["image_tp"], row["image_fp"], row["image_tn"], row["image_fn"], pct(row["image_fnr"]), pct(row["image_fpr"]), pct(row["image_specificity"])
        ])
    add_table(doc, ["Mô hình", "TP", "FP", "TN", "FN", "FNR", "FPR", "Specificity"], confusion_rows, [3.3, 1.3, 1.3, 1.3, 1.3, 2.1, 2.1, 3.3], caption="Bảng 5.3. Ma trận nhầm lẫn mức ảnh tại threshold đã khóa", font_size=9)
    add_para(doc, "SegFormer ưu tiên độ nhạy: 393/394 ảnh Defect được phát hiện nhưng 200/323 ảnh Good bị báo động. VMamba và U-Net đều bỏ sót 5 ảnh; VMamba giảm số báo động giả từ 105 xuống 70. Trong kiểm tra công nghiệp, chi phí của FN và FP không bằng nhau, vì vậy threshold cuối phải gắn với yêu cầu vận hành chứ không chỉ tối đa hóa Accuracy.")

    add_heading(doc, "5.4. Khả năng phát hiện theo kích thước (E3)", level=2)
    add_figure(doc, FIG_DIR / "06_e3_size_recall.png", "Hình 5.3. Recall theo kích thước vùng bất thường", width_cm=16.0)
    size_rows = []
    for size in ("Tiny", "Small", "Medium", "Large"):
        vm = next(r for r in by_size if r["model"] == "VMamba-T" and r["size_bin"] == size)
        size_rows.append([size, vm["n_regions"], pct(vm["region_recall"]), pct(vm["image_recall"]), pct(vm["positive_dice"]), pct(vm["positive_iou"])])
    add_table(doc, ["Nhóm", "Số vùng", "Recall vùng", "Recall ảnh", "Dice", "IoU"], size_rows, [2.5, 2.2, 2.8, 2.8, 2.8, 2.9], caption="Bảng 5.4. Kết quả VMamba-T theo kích thước bất thường", font_size=9.7)
    add_para(doc, "VMamba-T đạt region recall cao nhất ở nhóm Large (96,17%) và Dice tăng từ 74,64% ở Tiny lên 77,99% ở Large. Tiny vẫn đạt recall vùng 91,19%, cho thấy pipeline có khả năng phát hiện nhiều lỗi nhỏ; tuy nhiên đánh giá này phụ thuộc định nghĩa tứ phân vị suy ra từ Train và không thay thế phân tích theo kích thước vật lý thực tế.")

    add_heading(doc, "5.5. Phân tích theo nhóm khuyết tật (E4)", level=2)
    add_figure(doc, FIG_DIR / "08_e4_defect_group_performance.png", "Hình 5.4. Hiệu năng theo nhóm khuyết tật", width_cm=16.0)
    group_rows = []
    for group in ("bump", "bruise", "scratches", "Multiple-defects", "knife mark"):
        vm = next(r for r in by_group if r["model"] == "VMamba-T" and r["group"] == group)
        warning = "Có - mô tả" if vm["small_sample_warning"].lower() == "true" else "Không"
        group_rows.append([group, vm["n_images"], pct(vm["image_recall"]), pct(vm["positive_dice"]), pct(vm["positive_iou"]), warning])
    add_table(doc, ["Nhóm lỗi", "n", "Recall ảnh", "Dice", "IoU", "Cảnh báo"], group_rows, [3.5, 1.5, 2.7, 2.5, 2.5, 3.3], caption="Bảng 5.5. Kết quả VMamba-T theo nhóm lỗi", font_size=9.5)
    add_para(doc, "Nhóm Multiple-defects có Dice thấp nhất (40,10%) dù recall ảnh 100%, phản ánh việc nhận biết ảnh có lỗi dễ hơn định vị đầy đủ nhiều vùng. Knife mark có Dice 89,16% nhưng chỉ gồm 2 ảnh, do đó không thể suy rộng. Bump chiếm phần lớn Test và chi phối kết quả tổng; cần báo cáo theo nhóm để tránh kết luận bị che bởi mất cân bằng dữ liệu.")

    add_heading(doc, "5.6. Phân tích ảnh nhiều vùng lỗi (E5)", level=2)
    add_figure(doc, FIG_DIR / "10_e5_multi_region_performance.png", "Hình 5.5. Hiệu năng theo số lượng connected component", width_cm=15.7)
    mr_rows = []
    for label in ("single", "few", "many"):
        vm = next(r for r in multi_region if r["model"] == "VMamba-T" and r["multi_region_bin"] == label)
        mr_rows.append([label, vm["n_images"], vm["n_regions"], pct(vm["image_recall"]), pct(vm["region_recall"]), pct(vm["positive_dice"]), pct(vm["positive_iou"])])
    add_table(doc, ["Nhóm", "Ảnh", "Vùng", "Recall ảnh", "Recall vùng", "Dice", "IoU"], mr_rows, [2.2, 1.6, 1.8, 2.6, 2.7, 2.4, 2.7], caption="Bảng 5.6. Kết quả VMamba-T theo số vùng lỗi", font_size=9.3)
    add_para(doc, "VMamba-T đạt Dice cao nhất ở nhóm few (79,05%). Nhóm many có recall ảnh 100% nhưng recall vùng 90,88%, cho thấy mô hình thường nhận ra ảnh bất thường nhưng vẫn có thể bỏ một số vùng riêng lẻ. Đây là lý do cần đồng thời báo cáo image recall và region recall.")

    add_heading(doc, "5.7. Độ nhạy threshold (E7)", level=2)
    add_figure(doc, FIG_DIR / "11_e7_selected_thresholds.png", "Hình 5.6. Threshold được chọn trên Validation", width_cm=14.5)
    threshold_rows = [[r["model"], r["threshold"].replace(".", ","), "10%", "Đạt", "Tối đa Positive Dice, hòa chọn FPR thấp hơn"] for r in thresholds]
    add_table(doc, ["Mô hình", "Threshold", "Giới hạn FNR Val", "Trạng thái", "Quy tắc"], threshold_rows, [3.4, 2.2, 3.1, 2.0, 5.3], caption="Bảng 5.7. Kết quả lựa chọn threshold", font_size=9.2)
    add_para(doc, "Threshold khác nhau phản ánh mức hiệu chỉnh probability map của từng kiến trúc; không nên áp dụng chung 0,5 một cách mặc định. SegFormer cần threshold 0,66 nhưng vẫn có FPR cao trên Test, cho thấy ràng buộc FNR ưu tiên độ nhạy và có thể không đủ để kiểm soát báo động giả.")

    add_heading(doc, "5.8. Kết quả decision policy và cấu hình đề xuất", level=2)
    add_figure(doc, FIG_DIR / "12_decision_fpr_fnr_tradeoff.png", "Hình 5.7. Đánh đổi FPR-FNR của các chiến lược quyết định", width_cm=16.0)
    add_table(doc, ["Cấu hình", "FNR Test", "FPR Test", "Accuracy", "Review rate", "Lựa chọn"], [
        ["VMamba - threshold pixel gốc", "1,27%", "21,67%", "-", "0%", "Tham chiếu"],
        ["U-Net - threshold pixel gốc", "1,27%", "32,51%", "-", "0%", "Tham chiếu"],
        ["Spatial ensemble U-Net + VMamba", spatial_uv_fnr, spatial_uv_fpr, "-", spatial_uv_review, "Cấu hình vận hành mặc định"],
    ], [4.6, 2.3, 2.3, 2.3, 2.3, 4.2], caption="Bảng 5.8. So sánh cấu hình ngưỡng gốc và Spatial ensemble", font_size=9.2)
    add_para(doc, f"Spatial ensemble U-Net + VMamba đạt FNR {spatial_uv_fnr}, FPR {spatial_uv_fpr}, Positive Dice {spatial_uv_dice} và review rate {spatial_uv_review}. So với VMamba dùng threshold pixel gốc, policy giảm FPR khoảng {spatial_uv_fpr_reduction} điểm phần trăm nhưng FNR tăng khoảng {spatial_uv_fnr_increase} điểm. Vì vậy đây là lựa chọn cân bằng theo mục tiêu giảm báo động giả, không phải cấu hình tối đa hóa recall.")
    add_callout(doc, "Kết luận vận hành", "Nếu chi phí bỏ sót vượt xa chi phí kiểm tra lại, nên dùng chế độ safety triage có REVIEW hoặc hiệu chỉnh lại ràng buộc trên Validation. Không được hạ threshold dựa trên các ảnh Test đã biết.", "caution")

    add_heading(doc, "5.9. Phân tích định tính và lỗi còn lại (E8)", level=2)
    add_figure(doc, FIG_DIR / "20_qualitative_hybrid_um_errors.png", "Hình 5.8. Ví dụ báo động giả và bỏ sót của Spatial ensemble U-Net + VMamba", width_cm=15.8)
    add_para(doc, "Các báo động giả thường xuất hiện ở biên, lỗ gá, vùng phản xạ hoặc vệt có cấu trúc giống lỗi. Một số bỏ sót nằm ở vùng nhỏ, mảnh hoặc tương phản thấp. Kết quả định tính cho thấy decision policy đã giảm nhiều tín hiệu ngoài ROI nhưng chưa thay thế được việc làm sạch nhãn, bổ sung hard negatives và kiểm định trên dây chuyền thực.")

    add_heading(doc, "CHƯƠNG 6. TRIỂN KHAI, KIỂM THỬ VÀ ĐÁNH GIÁ", level=1, page_break=True)
    add_heading(doc, "6.1. Quy trình triển khai", level=2)
    add_bullets(doc, [
        "Đặt checkpoint cuối vào artifacts/checkpoints/final và policy vào artifacts/reports/final/decision_and_test_audit.",
        "Cài dependency backend; bảo đảm runtime VMamba tương thích Python, PyTorch, CUDA và compute capability.",
        "Chạy run_demo.ps1; kiểm tra /health trước khi suy luận.",
        "Frontend chỉ bật lựa chọn model nếu backend báo checkpoint và policy tương thích.",
        "Kết quả trả về gồm quyết định, lý do, số model đồng thuận, probability/mask và overlay từng model.",
    ], numbered=True)

    add_heading(doc, "6.2. Kết quả kiểm thử mới nhất", level=2)
    add_para(doc, "Bộ xác minh được chạy lại ngày 21/08/2026 trong workspace hiện tại. Protocol preflight xác nhận đúng 7.168 ảnh, không overlap và SHA-256 split khớp. Mười bốn kiểm thử Python PASS; frontend production build PASS; hai kiểm thử Node về server-rendered product và kết nối health/inference PASS. Launcher kết thúc với thông báo Submission verification: PASS.")
    add_table(doc, ["Hạng mục", "Số lượng", "Kết quả", "Nội dung kiểm tra"], [
        ["Protocol preflight", "1 quy trình", "PASS", "File, split overlap, content overlap, SHA-256"],
        ["Decision policy", "3 test", "PASS", "ROI biên tối; PASS/REVIEW/DEFECT; spatial consensus"],
        ["Training state", "4 test", "PASS", "Khôi phục best epoch và trạng thái"],
        ["Adaptive / Spatial ensemble", "3 test", "PASS", "ROI, connected components, đồng thuận theo vị trí và ràng buộc FNR"],
        ["Dataset Review Studio", "5 test", "PASS", "Mask, API, export không phá hủy, uncertain audit"],
        ["Web demo", "Build + 2 test", "PASS", "Render sản phẩm và wiring health/inference"],
    ], [4.2, 2.5, 2.0, 7.3], caption="Bảng 6.1. Tổng hợp kiểm thử ngày 21/08/2026", font_size=9.6)

    add_heading(doc, "6.3. Các nội dung chưa được xác minh", level=2)
    add_bullets(doc, [
        "VMamba selective-scan CUDA trên một môi trường Kaggle GPU mới.",
        "Full end-to-end API inference với ảnh thực sau lần đổi launcher cuối; health/config đã PASS nhưng phiên cleanup trước chưa chạy forward ảnh.",
        "Cài source release trên một máy hoặc thư mục sạch.",
        "Tái tạo toàn bộ ba quá trình huấn luyện từ đầu.",
        "Dataset review export round-trip sau lần cleanup hiện tại.",
        "Đánh giá hiện trường với điều kiện chiếu sáng, camera và phân bố lỗi của dây chuyền mục tiêu.",
    ])

    add_heading(doc, "6.4. Rủi ro và biện pháp kiểm soát", level=2)
    add_table(doc, ["Rủi ro", "Ảnh hưởng", "Kiểm soát hiện có", "Đề xuất tiếp theo"], [
        ["Rò rỉ Test", "Kết quả lạc quan giả", "Frozen split, hash, Validation-only policy", "Khóa artifact và reviewer độc lập"],
        ["Mất cân bằng nhóm lỗi", "Metric tổng che nhóm hiếm", "Báo cáo E4, cảnh báo n < 10", "Bổ sung dữ liệu knife mark/multiple"],
        ["Báo động giả biên/gá", "Tốn kiểm tra lại", "ROI biên tối, component policy", "Hard-negative mining theo dây chuyền"],
        ["Runtime VMamba đặc thù", "Không chạy trên máy mới", "Wheel và README môi trường", "Container/GPU compatibility matrix"],
        ["License dữ liệu chưa rõ", "Rủi ro phân phối", "Không commit dataset", "Xin xác nhận và chỉ phát hành manifest nếu cần"],
        ["Drift sản xuất", "Suy giảm recall/FPR", "Chưa có monitor hiện trường", "Log score, lấy mẫu review, tái hiệu chỉnh theo Validation mới"],
    ], [3.2, 3.2, 4.8, 4.8], caption="Bảng 6.2. Ma trận rủi ro kỹ thuật", font_size=8.8)

    add_heading(doc, "6.5. Đóng góp cá nhân và kỹ năng đạt được", level=2)
    add_callout(doc, "Cần cá nhân hóa", "Phần này phải được sinh viên chỉnh theo phân công thực tế và nhật ký làm việc; nội dung dưới đây là khung đối chiếu với sản phẩm trong repository.", "caution")
    add_bullets(doc, [
        "Khảo sát bài toán, thiết kế protocol và chuẩn hóa frozen split.",
        "Hiện thực/chuẩn hóa pipeline ba mô hình, full-resolution evaluation và threshold selection.",
        "Xây dựng decision policy Adaptive và Spatial ensemble, báo cáo kết quả và visualization.",
        "Phát triển Dataset Review Studio hoặc web demo theo phần việc thực tế.",
        "Viết test, tổ chức repository, tài liệu setup, log và hồ sơ bàn giao.",
        "Kỹ năng: PyTorch, xử lý ảnh, thiết kế thí nghiệm, đọc metric, API/web, kiểm thử và quản lý artifact.",
    ])

    add_heading(doc, "6.6. Sử dụng trí tuệ nhân tạo có trách nhiệm", level=2)
    add_para(doc, "Công cụ AI có thể được sử dụng để hỗ trợ rà soát cấu trúc, viết mã, tạo kiểm thử, phân tích log hoặc biên tập tài liệu. Tuy nhiên, mọi đề xuất phải được kiểm tra bằng nguồn gốc dữ liệu, mã chạy được, kiểm thử và artifact. Không dùng AI để tự điền thời gian thực tập, chữ ký, đóng góp cá nhân hoặc kết quả không tồn tại. Repository duy trì AI Development Log và prompt archive riêng với development log thông thường.")

    add_heading(doc, "KẾT LUẬN VÀ HƯỚNG PHÁT TRIỂN", level=1, page_break=True)
    add_heading(doc, "Kết luận", level=2)
    add_para(doc, "Đề tài đã xây dựng được một hệ thống tương đối hoàn chỉnh cho phân đoạn và hỗ trợ quyết định khuyết tật bề mặt nhôm: dữ liệu được audit và đóng băng; ba kiến trúc được huấn luyện, đánh giá theo cùng protocol; threshold và policy được chọn trên Validation; kết quả Test được báo cáo ở nhiều mức; ứng dụng review và demo được tích hợp; bộ kiểm thử xác minh các luồng quan trọng.")
    add_para(doc, f"Về câu hỏi RQ1, VMamba-T cho chất lượng mặt nạ và image F1 tốt nhất. Với RQ2, hiệu năng vẫn ổn định theo kích thước nhưng giảm rõ ở Multiple-defects và khi xét recall từng vùng của ảnh nhiều lỗi. Với RQ3, Spatial ensemble U-Net + VMamba giảm mạnh FPR so với threshold pixel gốc, đạt FPR {spatial_uv_fpr} và review rate {spatial_uv_review}, đổi lại FNR tăng lên {spatial_uv_fnr}. Với RQ4, hệ thống có cấu trúc bằng chứng và kiểm thử tốt, nhưng vẫn còn các bước xác minh môi trường sạch, CUDA và hiện trường.")
    add_heading(doc, "Hướng phát triển", level=2)
    add_bullets(doc, [
        "Thu thập thêm ảnh nhóm hiếm, đặc biệt knife mark và Multiple-defects; thực hiện review hai người cho ca bất đồng.",
        "Đánh giá calibration (ECE/Brier), bootstrap confidence interval và kiểm định thống kê giữa mô hình.",
        "Thực hiện E6 chính thức trên cùng GPU, cùng warm-up, batch và runtime; đo end-to-end latency gồm tiền/hậu xử lý.",
        "Đánh giá robustness theo ánh sáng, nhiễu, lệch camera và domain shift; xây dựng cơ chế monitor drift.",
        "Thử knowledge distillation hoặc model compression để giữ chất lượng VMamba với chi phí thấp hơn.",
        "Tách threshold/policy theo yêu cầu từng dây chuyền, nhưng chỉ hiệu chỉnh trên Validation đại diện và duy trì Test độc lập.",
        "Đóng gói container hoặc môi trường tái tạo; bổ sung clean-machine release test và hướng dẫn rollback.",
    ])

    add_heading(doc, "TÀI LIỆU THAM KHẢO", level=1, page_break=True)
    references = [
        "[1] E. Yang, P. Xing, H. Sun, W. Guo, Y. Ma, Z. Li, D. Zeng, “3CAD: A Large-Scale Real-World 3C Product Dataset for Unsupervised Anomaly Detection,” Proceedings of the AAAI Conference on Artificial Intelligence, vol. 39, no. 9, pp. 9175-9183, 2025. DOI: 10.1609/aaai.v39i9.32993.",
        "[2] O. Ronneberger, P. Fischer, T. Brox, “U-Net: Convolutional Networks for Biomedical Image Segmentation,” MICCAI, pp. 234-241, 2015. arXiv:1505.04597.",
        "[3] K. He, X. Zhang, S. Ren, J. Sun, “Deep Residual Learning for Image Recognition,” CVPR, pp. 770-778, 2016.",
        "[4] E. Xie, W. Wang, Z. Yu, A. Anandkumar, J. M. Alvarez, P. Luo, “SegFormer: Simple and Efficient Design for Semantic Segmentation with Transformers,” NeurIPS 34, 2021.",
        "[5] Y. Liu, Y. Tian, Y. Zhao, H. Yu, L. Xie, Y. Wang, Q. Ye, J. Jiao, Y. Liu, “VMamba: Visual State Space Model,” NeurIPS 37, 2024. DOI: 10.52202/079017-3273.",
        "[6] I. Loshchilov, F. Hutter, “Decoupled Weight Decay Regularization,” ICLR, 2019.",
        "[7] L. R. Dice, “Measures of the Amount of Ecologic Association Between Species,” Ecology, vol. 26, no. 3, pp. 297-302, 1945.",
        "[8] T. Fawcett, “An Introduction to ROC Analysis,” Pattern Recognition Letters, vol. 27, no. 8, pp. 861-874, 2006.",
        "[9] PyTorch Contributors, “PyTorch Documentation.” https://pytorch.org/docs/ (truy cập 21/08/2026).",
        "[10] FastAPI, “FastAPI Documentation.” https://fastapi.tiangolo.com/ (truy cập 21/08/2026).",
        "[11] React Team, “React Documentation.” https://react.dev/ (truy cập 21/08/2026).",
        "[12] Dự án 3CAD chính thức, https://github.com/EnquanYang2022/3CAD (truy cập 21/08/2026).",
    ]
    for ref in references:
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(-0.8)
        p.paragraph_format.left_indent = Cm(0.8)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.15
        r = p.add_run(ref)
        set_font(r, size=11.5, color=INK)

    add_heading(doc, "PHỤ LỤC A. HƯỚNG DẪN CHẠY NHANH", level=1, page_break=True)
    add_para(doc, "Các lệnh sau được chạy tại thư mục gốc dự án trên Windows PowerShell.")
    command_rows = [
        ["Xác minh toàn bộ", ".\\verify.ps1 -IncludeWeb"],
        ["Mở Dataset Review Studio", ".\\run_review.ps1"],
        ["Mở web demo", ".\\run_demo.ps1"],
        ["Chạy lại thí nghiệm ba model từ cache", ".\\scripts\\experiments\\run_three_model_experiments.ps1"],
        ["Retrain U-Net", ".\\scripts\\training\\retrain_models.ps1 -Model unet -RunName <ten_run>"],
    ]
    add_table(doc, ["Mục đích", "Lệnh"], command_rows, [5.3, 10.7], caption="Bảng A.1. Các launcher chính", font_size=10)
    add_callout(doc, "Yêu cầu", "VMamba cần wheel và CUDA tương thích. Khi thiếu checkpoint, ứng dụng phải báo unavailable; không được tạo dự báo thay thế.")

    add_heading(doc, "PHỤ LỤC B. ARTIFACT VÀ BẰNG CHỨNG", level=1, page_break=True)
    add_table(doc, ["Artifact", "Vị trí", "Mục đích"], [
        ["Checkpoint cuối", "artifacts/checkpoints/final/", "Ba best checkpoint"],
        ["Bảng E2-E8", "artifacts/reports/final/thesis_evaluation_report/", "Kết quả luận giải khoa học"],
        ["Decision audit", "artifacts/reports/final/decision_and_test_audit/", "Policy, OOF, Test và test-case audit"],
        ["Visualizations", "artifacts/reports/final/visualizations/", "20 PNG, SVG, dashboard"],
        ["Verification report", "submission/07_evidence/VERIFICATION_REPORT.md", "Giao thức và test"],
        ["Dataset provenance", "submission/07_evidence/DATASET_PROVENANCE.md", "Nguồn, số lượng và license"],
        ["Model manifest", "submission/07_evidence/MODEL_ARTIFACT_MANIFEST.csv", "Hash, epoch, threshold"],
        ["Kaggle training outputs", "artifacts/training_download/<model>/...", "History, summary, learning curves, environment và best.pt của U-Net, SegFormer, VMamba"],
        ["AI log", "submission/05_logs/AI_DEVELOPMENT_LOG.md", "Theo dõi sử dụng AI"],
    ], [3.5, 7.4, 5.1], caption="Bảng B.1. Danh mục bằng chứng quan trọng", font_size=9.2)

    add_heading(doc, "PHỤ LỤC C. CHECKLIST HOÀN THIỆN BẢN NỘP", level=1, page_break=True)
    add_bullets(doc, [
        "Đã thay toàn bộ trường [ĐIỀN ...] và bỏ bôi vàng trên trang bìa.",
        "Đã chuyển nội dung vào template chính thức của khoa hoặc xác nhận định dạng hiện tại được chấp nhận.",
        "Đã cập nhật mục lục và số trang bằng Ctrl+A, F9.",
        "Đã cá nhân hóa lời cảm ơn, lời cam đoan, đóng góp cá nhân và nhật ký thực tập.",
        "Đã bổ sung nhận xét/chữ ký/xác nhận của đơn vị theo biểu mẫu bắt buộc.",
        "Đã chạy verify.ps1 -IncludeWeb và lưu log/commit hash của lần cuối.",
        "Đã kiểm tra mọi số liệu trong báo cáo khớp artifact đóng băng.",
        "Đã xác minh quyền sử dụng ảnh trước khi nộp hoặc công bố dataset kèm theo.",
        "Đã xuất PDF và kiểm tra trực quan toàn bộ trang, hình, bảng, header/footer.",
    ])
    add_callout(doc, "Trạng thái bản này", "Nội dung kỹ thuật đã được tổng hợp từ workspace và số liệu artifact. Thông tin cá nhân, đơn vị, thời gian thực tập, chữ ký và yêu cầu biểu mẫu chính thức vẫn thuộc trách nhiệm người nộp.", "caution")

    props = doc.core_properties
    props.title = "Báo cáo Thực tập tốt nghiệp - Phân đoạn khuyết tật bề mặt nhôm"
    props.subject = "Hệ thống phân đoạn và hỗ trợ quyết định dựa trên U-Net, SegFormer và VMamba"
    props.author = "[ĐIỀN HỌ VÀ TÊN SINH VIÊN]"
    props.keywords = "3CAD, ANI, aluminum defect segmentation, U-Net, SegFormer, VMamba"
    props.comments = "Bản kỹ thuật được tạo từ artifact dự án; cần điền thông tin cá nhân và đối chiếu template khoa trước khi nộp."

    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    print(build_report())
